import docx

doc_path = r"C:\Users\ben.lefort\OneDrive - Ontario Federation of Agriculture\Desktop\Municipal Bylaw Database\Data Pulls\Reports\Province Wide\All bylaws\county-bylaw-maps\How to use database\OFA Bylaw Database - User Onboarding Guide.docx"
doc = docx.Document(doc_path)

found_idx = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith("1. Choose your keywords."):
        found_idx = i
        break

if found_idx != -1:
    p = doc.paragraphs[found_idx]
    run = p.add_run("\n   • Negative Keywords (Optional): If you want to exclude false positives (e.g., searching for \"greenhouse\" but want to ignore \"greenhouse gas\"), type those exact phrases into the Negative Keywords box. The scanner will cleanly replace them with [IGNORED] so they don't falsely trigger a match.")
    doc.save(doc_path)
    print("Successfully updated the document.")
else:
    print("Could not find the target paragraph.")
