import docx
import os

doc_path = r"C:\Users\ben.lefort\OneDrive - Ontario Federation of Agriculture\Desktop\Municipal Bylaw Database\Data Pulls\Reports\Province Wide\All bylaws\county-bylaw-maps\How to use database\OFA Bylaw Database - User Onboarding Guide.docx"
doc = docx.Document(doc_path)

# Verify we are on the right paragraph
if doc.paragraphs[183].text.startswith("2. Choose which municipalities to scan"):
    doc.paragraphs[183].text = '2. Choose which municipalities to scan. You can choose "Select All — Province-Wide Scan" to scan all 444 municipalities (this takes 3-5 minutes), choose "Select by Region/County" to select one or more upper-tier regions (the scanner will automatically target all lower-tier municipalities within those regions), or choose "Select Specific Municipalities" to pick them individually from the dropdown.'
    doc.save(doc_path)
    print("Document successfully updated.")
else:
    print("Error: Paragraph 183 does not match expected text.")
    print("Found:", doc.paragraphs[183].text)
