"""
report_generator.py — Generates OFA-branded DOCX bylaw reports.

Supports three scopes:
  - Provincial: Province-wide overview with county-level aggregations
  - County: Full detail for all municipalities in one county/region
  - Municipality: Deep dive on a single municipality
"""

import io
import os
import datetime
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from db_utils import get_connection, get_report_data, resolve_yes_no

# ── Paths ──
HERE = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(HERE, "OFA_logo.png")

# ── OFA Branding ──
OFA_RED = RGBColor(0xCC, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEADER_BG = "333333"
GREEN_BG = "D4EDDA"
RED_BG = "F8D7DA"
AMBER_BG = "FFF3CD"
LIGHT_GRAY_BG = "F2F2F2"

CATEGORY_ORDER = ["DC", "STORMWATER", "SITE_ALT", "LGD", "TREES", "CHICKENS", "FENCES"]

CATEGORY_NAMES = {
    "DC": "Development Charges",
    "STORMWATER": "Stormwater",
    "SITE_ALT": "Site Alteration & Fill",
    "LGD": "Livestock Guardian Dogs",
    "TREES": "Tree / Forest Conservation",
    "CHICKENS": "Backyard Chickens",
    "FENCES": "Fence Bylaws",
}

DETAIL_FIELDS = {
    "DC": [
        ("has_dc", "Has DC"),
        ("fees_bylaw_name", "Fees Bylaw"),
        ("fees_enacted", "Fees Enacted"),
        ("fees_expiry", "Fees Expiry"),
        ("fees_expiry_notes", "Fees Notes"),
    ],
    "STORMWATER": [
        ("charge_type", "Charge Type"),
        ("fee_calculation", "Fee Calculation"),
    ],
    "SITE_ALT": [
        ("farm_exemption", "Farm Exemption"),
        ("special_provision", "Special Provision"),
        ("guidelines_wording", "Guidelines"),
        ("exception_wording", "Exception Wording"),
    ],
    "LGD": [
        ("has_lgd_definition", "Has LGD Def"),
        ("lgd_definition", "LGD Definition"),
        ("has_herding_def", "Has Herding Def"),
        ("herding_definition", "Herding Definition"),
        ("exempt_license_fees", "Exempt License Fees"),
        ("collar_tag_req", "Collar/Tag Req"),
        ("barking_restrictions", "Barking Restrictions"),
        ("exempt_barking", "Exempt Barking"),
        ("dog_limit", "Dog Limit"),
    ],
    "TREES": [
        ("farm_exemption", "Farm Exemption"),
        ("bylaw_wording", "Bylaw Wording"),
        ("farming_exception_wording", "Exception Wording"),
        ("list_of_exceptions", "Exceptions List"),
    ],
    "CHICKENS": [
        ("can_keep", "Can Keep"),
        ("definition", "Definition"),
        ("chicken_limit", "Limit"),
        ("roosters_allowed", "Roosters"),
        ("licence_required", "Licence Req"),
        ("welfare_requirements", "Welfare Req"),
    ],
    "FENCES": [
        ("has_fence_bylaw", "Has Bylaw"),
        ("applies_all_lands", "All Lands"),
        ("replaces_lfa", "Replaces LFA"),
        ("security_fencing_exemption", "Security Fence Exempt"),
        ("electrified_fencing_exemption", "Electric Fence Exempt"),
        ("equal_apportionment", "Equal Apportionment"),
        ("fence_notes", "Notes"),
    ],
}


# ══════════════════════════════════════════════════════════════
# Formatting helpers
# ══════════════════════════════════════════════════════════════

def _clean(val, max_len=150):
    """Clean a value for display: handle None/nan, truncate long text."""
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return "—"
    s = str(val).strip()
    if not s or s.lower() == "nan":
        return "—"
    if len(s) > max_len:
        s = s[:max_len - 3] + "..."
    return s


def _resolve(val):
    """Resolve integer-coded Yes/No and clean."""
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return "—"
    s = str(val).strip()
    if s in ("1", "2", "3", "4", "5"):
        return resolve_yes_no(val)
    return resolve_yes_no(val) if s else "—"


def _set_cell_shading(cell, hex_color):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _status_bg(status):
    """Return background hex for an exemption status string."""
    s = str(status).strip().upper()
    if s == "YES":
        return GREEN_BG
    if s == "NO":
        return RED_BG
    if s in ("N/A", "NOT KNOWN", "NO EXPLICIT EXEMPTION FOUND"):
        return AMBER_BG
    return None


def _set_cell(cell, text, bold=False, size=Pt(9), color=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = size
    run.font.name = "Calibri"
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color


def _add_table(doc, headers, rows):
    """Create a styled table with header row and alternating shading."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell(cell, h, bold=True, size=Pt(8), color=WHITE)
        _set_cell_shading(cell, HEADER_BG)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        for c_idx, val in enumerate(row_data):
            _set_cell(row.cells[c_idx], _clean(val, 200), size=Pt(8))
        if r_idx % 2 == 1:
            for cell in row.cells:
                _set_cell_shading(cell, LIGHT_GRAY_BG)

    return table


def _heading(doc, text, level=1):
    """Add headed with OFA red."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = OFA_RED
    return h


# ══════════════════════════════════════════════════════════════
# Report sections
# ══════════════════════════════════════════════════════════════

def _add_cover_page(doc, scope, scope_value):
    """OFA-branded cover page."""
    for _ in range(3):
        doc.add_paragraph("")

    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(LOGO_PATH, width=Inches(3))

    doc.add_paragraph("")

    # Title
    titles = {
        'provincial': ("Ontario Municipal Bylaw Report", "Province-Wide Overview"),
        'county': (f"{scope_value}\nMunicipal Bylaw Report", "County / Region Overview"),
        'municipality': (f"{scope_value}\nBylaw Report", "Municipality Report"),
    }
    title, subtitle = titles.get(scope, ("Bylaw Report", ""))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = OFA_RED
    run.bold = True
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_GRAY
    run.font.name = "Calibri"

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared by the Ontario Federation of Agriculture")
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = DARK_GRAY

    doc.add_page_break()


def _add_executive_summary(doc, data):
    """Exemption scorecard across all 7 categories."""
    munis = data['municipalities']
    bylaws = data['bylaws']
    signals = data['signals']
    scope = data['scope']

    _heading(doc, "Executive Summary")

    n_munis = len(munis)
    n_signals = len(signals)

    if scope == 'provincial':
        n_counties = munis['geographic_area'].nunique()
        doc.add_paragraph(
            f"This report covers {n_munis} municipalities across {n_counties} "
            f"counties and regions in Ontario, with {n_signals} recent intelligence signals."
        )
    elif scope == 'county':
        doc.add_paragraph(
            f"This report covers {n_munis} municipalities in {data['scope_value']}, "
            f"with {n_signals} recent intelligence signals for this area."
        )
    else:
        doc.add_paragraph(
            f"Detailed bylaw analysis for {data['scope_value']}."
        )

    # Scorecard table
    _heading(doc, "Exemption Scorecard", level=2)
    headers = ["Category", "Yes", "No", "N/A", "Unknown", "Total", "Bylaw %", "Exempt %"]
    rows = []

    for cat in CATEGORY_ORDER:
        cat_bylaws = bylaws[bylaws['category'] == cat]
        if cat_bylaws.empty:
            rows.append([CATEGORY_NAMES[cat], "0", "0", "0", "0", "0", "—", "—"])
            continue
        statuses = cat_bylaws['exemption_status'].apply(_resolve)
        yes = (statuses == "Yes").sum()
        no = (statuses == "No").sum()
        na = statuses.isin(["N/A"]).sum()
        nk = len(statuses) - yes - no - na
        total = len(statuses)
        
        active_bylaws = yes + no
        bylaw_cov = f"{active_bylaws / total * 100:.0f}%" if total > 0 else "—"
        exempt_cov = f"{yes / active_bylaws * 100:.0f}%" if active_bylaws > 0 else "—"
        
        rows.append([CATEGORY_NAMES[cat], str(yes), str(no), str(na), str(nk), str(total), bylaw_cov, exempt_cov])

    table = _add_table(doc, headers, rows)

    # Color the Exempt % column
    for i, row_data in enumerate(rows):
        cov = row_data[7]
        if cov != "—":
            pct = int(cov.replace("%", ""))
            bg = GREEN_BG if pct >= 70 else AMBER_BG if pct >= 40 else RED_BG
            _set_cell_shading(table.rows[i + 1].cells[7], bg)

    doc.add_page_break()


def _add_category_section(doc, data, cat_code):
    """One section per bylaw category."""
    scope = data['scope']
    bylaws = data['bylaws']
    details_df = data['details'].get(cat_code, pd.DataFrame())
    cat_name = CATEGORY_NAMES[cat_code]
    cat_bylaws = bylaws[bylaws['category'] == cat_code].copy()

    _heading(doc, cat_name)

    if cat_bylaws.empty:
        doc.add_paragraph("No data available for this category.")
        doc.add_page_break()
        return

    # Summary stats
    statuses = cat_bylaws['exemption_status'].apply(_resolve)
    yes = (statuses == "Yes").sum()
    no = (statuses == "No").sum()
    na = statuses.isin(["N/A"]).sum()
    total = len(statuses)
    
    active_bylaws = yes + no
    bylaw_cov = f"{active_bylaws / total * 100:.0f}%" if total > 0 else "0%"
    exempt_cov = f"{yes / active_bylaws * 100:.0f}%" if active_bylaws > 0 else "N/A"
    
    doc.add_paragraph(
        f"Bylaw presence: {bylaw_cov} ({active_bylaws} active bylaws identified out of {total} municipalities)."
    )
    doc.add_paragraph(
        f"Farm exemption coverage: {exempt_cov} ({yes} out of {active_bylaws} active bylaws). "
        f"{no} active bylaws have no exemption, and {na} municipalities explicitly do not have this type of bylaw."
    )

    if scope == 'provincial':
        _add_category_provincial(doc, cat_bylaws)
    else:
        _add_category_detail(doc, cat_code, cat_bylaws, details_df)

    doc.add_page_break()


def _add_category_provincial(doc, cat_bylaws):
    """County-level aggregation table for provincial reports."""
    _heading(doc, "County / Region Summary", level=2)

    cat_bylaws = cat_bylaws.copy()
    cat_bylaws['_status'] = cat_bylaws['exemption_status'].apply(_resolve)

    headers = ["County / Region", "Yes", "No", "N/A", "Unknown", "Total", "Bylaw %", "Exempt %"]
    rows = []

    for county, group in sorted(cat_bylaws.groupby('geographic_area'), key=lambda x: str(x[0])):
        county = county if county and str(county).strip() else "Unassigned"
        s = group['_status']
        yes = (s == "Yes").sum()
        no = (s == "No").sum()
        na = s.isin(["N/A"]).sum()
        nk = len(s) - yes - no - na
        total = len(s)
        
        active_bylaws = yes + no
        bylaw_cov = f"{active_bylaws / total * 100:.0f}%" if total > 0 else "—"
        exempt_cov = f"{yes / active_bylaws * 100:.0f}%" if active_bylaws > 0 else "—"
        
        rows.append([county, str(yes), str(no), str(na), str(nk), str(total), bylaw_cov, exempt_cov])

    _add_table(doc, headers, rows)


def _add_category_detail(doc, cat_code, cat_bylaws, details_df):
    """Municipality-level table with full detail sub-fields."""
    _heading(doc, "Municipality Overview", level=2)

    headers = ["Municipality", "Exemption", "Bylaw Name", "Enacted", "Expiry"]
    rows = []
    for _, b in cat_bylaws.iterrows():
        status = _resolve(b['exemption_status'])
        rows.append([
            _clean(b.get('municipality_name')),
            status,
            _clean(b.get('bylaw_name')),
            _clean(b.get('date_enacted')),
            _clean(b.get('expiry_date')),
        ])

    table = _add_table(doc, headers, rows)
    for i, row_data in enumerate(rows):
        bg = _status_bg(row_data[1])
        if bg:
            _set_cell_shading(table.rows[i + 1].cells[1], bg)

    # Detail sub-fields
    detail_fields = DETAIL_FIELDS.get(cat_code, [])
    if not detail_fields or details_df.empty:
        return

    merged = cat_bylaws.merge(
        details_df.drop(columns=['id'], errors='ignore'),
        left_on='id', right_on='bylaw_id', how='inner', suffixes=('', '_d'))

    if merged.empty:
        return

    # Keep rows that have at least one non-empty detail field
    def _has_data(row):
        return any(
            str(row.get(f, '') or '').strip() not in ('', 'None', 'nan')
            for f, _ in detail_fields
        )

    merged = merged[merged.apply(_has_data, axis=1)]
    if merged.empty:
        return

    doc.add_paragraph("")
    _heading(doc, "Detailed Information", level=2)

    d_headers = ["Municipality"] + [display for _, display in detail_fields]
    d_rows = []
    for _, row in merged.iterrows():
        d_row = [_clean(row.get('municipality_name'))]
        for field, _ in detail_fields:
            val = row.get(field, '')
            val = _resolve(val) if str(val).strip() in ('1', '2', '3', '4', '5') else _clean(val)
            d_row.append(val)
        d_rows.append(d_row)

    _add_table(doc, d_headers, d_rows)


def _add_municipality_profiles(doc, data):
    """Individual profiles for each municipality (county scope)."""
    munis = data['municipalities']
    bylaws = data['bylaws']
    contacts = data['contacts']
    signals = data['signals']

    _heading(doc, "Municipality Profiles")

    for _, muni in munis.iterrows():
        muni_id = muni['id']
        muni_name = muni['name']

        _heading(doc, muni_name, level=2)

        # Info line
        parts = []
        if muni.get('municipal_status'):
            parts.append(f"Status: {muni['municipal_status']}")
        if muni.get('geographic_area'):
            parts.append(f"Region: {muni['geographic_area']}")
        if muni.get('website'):
            parts.append(f"Website: {muni['website']}")
        if parts:
            doc.add_paragraph(" | ".join(parts))

        # Contact info
        contact_parts = []
        muni_contacts = contacts[contacts['municipality_id'] == muni_id]
        if not muni_contacts.empty:
            c = muni_contacts.iloc[0]
            name = f"{c.get('first_name', '') or ''} {c.get('last_name', '') or ''}".strip()
            if name:
                contact_parts.append(f"Contact: {name}")
            if c.get('email'):
                contact_parts.append(f"Email: {c['email']}")
            if c.get('phone'):
                contact_parts.append(f"Phone: {c['phone']}")
        else:
            if muni.get('contact_name'):
                contact_parts.append(f"Contact: {muni['contact_name']}")
            if muni.get('clerk_email'):
                contact_parts.append(f"Email: {muni['clerk_email']}")
            if muni.get('clerk_phone'):
                contact_parts.append(f"Phone: {muni['clerk_phone']}")
        if contact_parts:
            doc.add_paragraph(" | ".join(contact_parts))

        # Bylaw summary table
        muni_bylaws = bylaws[bylaws['municipality_id'] == muni_id]
        if not muni_bylaws.empty:
            headers = ["Category", "Exemption", "Bylaw", "Expiry"]
            rows = []
            for _, b in muni_bylaws.iterrows():
                status = _resolve(b['exemption_status'])
                rows.append([
                    CATEGORY_NAMES.get(b['category'], b['category']),
                    status,
                    _clean(b.get('bylaw_name')),
                    _clean(b.get('expiry_date')),
                ])
            table = _add_table(doc, headers, rows)
            for i, rd in enumerate(rows):
                bg = _status_bg(rd[1])
                if bg:
                    _set_cell_shading(table.rows[i + 1].cells[1], bg)

        # Signals for this muni
        muni_signals = signals[signals['municipality_id'] == muni_id]
        if not muni_signals.empty:
            p = doc.add_paragraph()
            run = p.add_run("Intelligence Alert: ")
            run.bold = True
            run.font.color.rgb = OFA_RED
            run.font.size = Pt(10)
            for _, sig in muni_signals.iterrows():
                doc.add_paragraph(
                    f"{_clean(sig.get('trigger_keyword'))} ({sig.get('category', '')}) "
                    f"— {_clean(sig.get('discovered_date'))}"
                )

        doc.add_paragraph("")  # spacing


def _add_intelligence_section(doc, data):
    """Scanner intelligence alerts."""
    signals = data['signals']

    _heading(doc, "Intelligence Alerts")

    if signals.empty:
        doc.add_paragraph("No recent scanner signals detected.")
        doc.add_page_break()
        return

    doc.add_paragraph(
        f"{len(signals)} scanner signal(s) detected. These are municipalities where "
        f"council agendas or minutes mentioned OFA-tracked bylaw keywords."
    )

    headers = ["Municipality", "Category", "Keyword", "Date", "Summary / Snippet"]
    rows = []
    for _, sig in signals.iterrows():
        ai = _clean(sig.get('ai_summary'), 120)
        if ai == "—":
            ai = _clean(sig.get('snippet'), 120)
        rows.append([
            _clean(sig.get('municipality_name', sig.get('munid_raw'))),
            CATEGORY_NAMES.get(sig.get('category', ''), _clean(sig.get('category'))),
            _clean(sig.get('trigger_keyword')),
            _clean(sig.get('discovered_date')),
            ai,
        ])

    _add_table(doc, headers, rows)
    doc.add_page_break()


def _add_advocacy_matrix(doc, data):
    """Priority ranking with auto-generated recommended actions."""
    bylaws = data['bylaws']
    signals = data['signals']
    munis = data['municipalities']
    scope = data['scope']

    _heading(doc, "Advocacy Priority Matrix")
    doc.add_paragraph(
        "Municipalities ranked by advocacy urgency. Scoring: missing exemptions (3 pts), "
        "expired bylaws (5 pts), expiring within 18 months (3 pts), scanner signals (2 pts), "
        "unknown status (1 pt). Recommended actions are auto-generated."
    )

    today = datetime.date.today()
    scored = []

    for _, muni in munis.iterrows():
        muni_id = muni['id']
        muni_name = muni['name']
        mb = bylaws[bylaws['municipality_id'] == muni_id]
        ms = signals[signals['municipality_id'] == muni_id] if not signals.empty else pd.DataFrame()

        score = 0
        no_cats, expired_cats, expiring_cats, unknown_cats = [], [], [], []

        for _, b in mb.iterrows():
            status = _resolve(b['exemption_status'])
            cat_label = CATEGORY_NAMES.get(b['category'], b['category'])

            if status == "No":
                no_cats.append(cat_label)
                score += 3
            elif status in ("NOT KNOWN", "No explicit exemption found", "—"):
                unknown_cats.append(cat_label)
                score += 1

            expiry = str(b.get('expiry_date', '') or '').strip()
            if expiry:
                try:
                    exp_date = pd.to_datetime(expiry).date()
                    if exp_date < today:
                        expired_cats.append(cat_label)
                        score += 5
                    elif (exp_date - today).days <= 540:
                        expiring_cats.append(cat_label)
                        score += 3
                except Exception:
                    pass

        if not ms.empty:
            score += 2 * len(ms)

        if score == 0:
            continue

        # Recommended action (highest priority first)
        if expired_cats:
            action = f"Monitor renewal — request exemption ({', '.join(expired_cats[:2])})"
        elif no_cats:
            action = f"Send template letter ({', '.join(no_cats[:2])})"
        elif expiring_cats:
            action = f"Engage before renewal ({', '.join(expiring_cats[:2])})"
        elif not ms.empty:
            kw = _clean(ms.iloc[0].get('trigger_keyword', 'activity'), 30)
            action = f"Review council activity — '{kw}'"
        elif unknown_cats:
            action = f"Research needed ({', '.join(unknown_cats[:2])})"
        else:
            action = "Monitor"

        scored.append({
            'name': muni_name, 'score': score,
            'no': len(no_cats), 'expired': len(expired_cats),
            'expiring': len(expiring_cats), 'signals': len(ms),
            'action': action,
        })

    if not scored:
        doc.add_paragraph("No advocacy priorities identified.")
        return

    scored.sort(key=lambda x: x['score'], reverse=True)

    if scope == 'provincial':
        scored = scored[:30]
        doc.add_paragraph(f"Showing top 30 priority municipalities.")

    headers = ["Municipality", "Priority", "No Exempt.", "Expired", "Expiring", "Signals", "Recommended Action"]
    rows = []
    for s in scored:
        pri = "Critical" if s['score'] >= 10 else "High" if s['score'] >= 5 else "Monitor"
        rows.append([
            s['name'], pri,
            str(s['no']) if s['no'] else "—",
            str(s['expired']) if s['expired'] else "—",
            str(s['expiring']) if s['expiring'] else "—",
            str(s['signals']) if s['signals'] else "—",
            s['action'],
        ])

    table = _add_table(doc, headers, rows)

    for i, s in enumerate(scored):
        bg = RED_BG if s['score'] >= 10 else AMBER_BG if s['score'] >= 5 else GREEN_BG
        _set_cell_shading(table.rows[i + 1].cells[1], bg)

    doc.add_page_break()


def _add_methodology(doc, data):
    """Data sources and definitions."""
    _heading(doc, "Methodology & Data Notes")

    doc.add_paragraph(
        f"Generated on {datetime.date.today().strftime('%B %d, %Y')} from the "
        f"Ontario Federation of Agriculture Municipal Bylaw Database."
    )

    _heading(doc, "Categories Tracked", level=2)
    for cat in CATEGORY_ORDER:
        doc.add_paragraph(f"{CATEGORY_NAMES[cat]}", style='List Bullet')

    _heading(doc, "Exemption Status Definitions", level=2)
    for status, desc in [
        ("Yes", "An explicit farm or agricultural exemption exists in the bylaw."),
        ("No", "The bylaw applies to farms with no exemption provision."),
        ("N/A", "The municipality does not have this type of bylaw."),
        ("NOT KNOWN", "Research has not yet confirmed the exemption status."),
    ]:
        doc.add_paragraph(f"{status}: {desc}", style='List Bullet')

    _heading(doc, "Priority Scoring", level=2)
    for item in [
        "Missing exemption (No): 3 points per category",
        "Expired bylaw: 5 points per category",
        "Expiring within 18 months: 3 points per category",
        "Scanner signal detected: 2 points per signal",
        "Unknown status: 1 point per category",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Disclaimer: ")
    run.bold = True
    p.add_run(
        "This report is for informational purposes. Bylaw data may change between "
        "database updates. Always consult the original municipal bylaw documents."
    )


# ══════════════════════════════════════════════════════════════
# Main entry point
# ══════════════════════════════════════════════════════════════

def generate_report(scope='provincial', scope_value=None):
    """Generate a complete DOCX report.

    Args:
        scope: 'provincial', 'county', or 'municipality'
        scope_value: County name or municipality name

    Returns:
        BytesIO containing the generated DOCX file.
    """
    conn = get_connection()
    data = get_report_data(conn, scope, scope_value)
    conn.close()

    doc = Document()

    # Global font defaults
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    _add_cover_page(doc, scope, scope_value)
    _add_executive_summary(doc, data)

    for cat_code in CATEGORY_ORDER:
        _add_category_section(doc, data, cat_code)

    if scope == 'county':
        _add_municipality_profiles(doc, data)

    _add_intelligence_section(doc, data)
    _add_advocacy_matrix(doc, data)
    _add_methodology(doc, data)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
