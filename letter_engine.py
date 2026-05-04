"""
letter_engine.py — Personalized advocacy letter generation and mailto delivery.

Handles:
  - Resolving the best generic recipient email for a municipality
  - Filling [INSERT ...] placeholders in OFA template .docx letters
  - Rendering letters as plain text for mailto: body
  - Generating mailto: links with pre-filled recipient, subject, and body

Used by the Lower Tier and Upper Tier map pages in the Streamlit app.
"""

import io
import os
import re
import copy
import datetime
import urllib.parse

from docx import Document

# ── Paths ──
HERE = os.path.dirname(os.path.abspath(__file__))


# ══════════════════════════════════════════════════════════════
# Recipient Resolution
# ══════════════════════════════════════════════════════════════

def get_recipient_email(conn, municipality_id, municipality_name=None):
    """Resolve the best generic recipient email for a municipality.

    Priority chain:
      1. contacts.municipal_email (generic inbox — info@, clerk@, admin@)
      2. municipalities.clerk_email (sparse fallback)
      3. None (user must enter manually)

    Args:
        conn: Database connection (PgWrapper or sqlite3 connection)
        municipality_id: Integer ID of the municipality
        municipality_name: Optional name for display

    Returns:
        dict with keys: email, source, contact_name, municipality_name
              email may be None if no address found
    """
    result = {
        "email": None,
        "source": None,
        "contact_name": None,
        "municipality_name": municipality_name or "",
    }

    # Use PgWrapper.execute() if available (handles ? → %s conversion),
    # otherwise fall back to raw cursor for sqlite3
    def _execute(query, params):
        if hasattr(conn, "execute"):
            return conn.execute(query, params)
        raw = getattr(conn, "conn", conn)
        cur = raw.cursor()
        cur.execute(query, params)
        return cur

    # Default address fetch (applies to both Priority 1 and 2)
    address_str = ""
    try:
        cur_addr = _execute("""
            SELECT address_1, address_2, city, postal
            FROM contacts
            WHERE municipality_id = ?
              AND address_1 IS NOT NULL
            LIMIT 1
        """, (municipality_id,))
        row_addr = cur_addr.fetchone()
        if row_addr:
            a1, a2, city, postal = row_addr[0] or "", row_addr[1] or "", row_addr[2] or "", row_addr[3] or ""
            parts = [a1]
            if a2: parts.append(a2)
            if city: parts.append(city)
            if postal: parts.append(postal)
            address_str = ", ".join(parts)
    except Exception:
        pass
        
    result["address"] = address_str

    # ── Priority 1: contacts.municipal_email ──
    try:
        cur = _execute("""
            SELECT c.municipal_email, c.first_name, c.last_name, c.title
            FROM contacts c
            WHERE c.municipality_id = ?
              AND c.municipal_email IS NOT NULL
              AND length(c.municipal_email) > 2
            ORDER BY
                CASE WHEN c.title LIKE '%%Clerk%%' AND c.title NOT LIKE '%%Deputy%%' THEN 0
                     WHEN c.title LIKE '%%Clerk%%' THEN 1
                     WHEN c.title LIKE '%%CAO%%' THEN 2
                     ELSE 3
                END
            LIMIT 1
        """, (municipality_id,))
        row = cur.fetchone()
        if row:
            email = row[0] if isinstance(row, (list, tuple)) else row["municipal_email"]
            result["email"] = email
            result["source"] = "municipal_email"
            if isinstance(row, (list, tuple)):
                fname, lname = row[1] or "", row[2] or ""
            else:
                fname, lname = row.get("first_name", "") or "", row.get("last_name", "") or ""
            name = f"{fname} {lname}".strip()
            result["contact_name"] = name if name else None
            return result
    except Exception:
        pass

    # ── Priority 2: municipalities.clerk_email ──
    try:
        cur = _execute("""
            SELECT clerk_email, contact_name
            FROM municipalities
            WHERE id = ?
              AND clerk_email IS NOT NULL
              AND length(clerk_email) > 2
        """, (municipality_id,))
        row = cur.fetchone()
        if row:
            email = row[0] if isinstance(row, (list, tuple)) else row["clerk_email"]
            result["email"] = email
            result["source"] = "clerk_email"
            contact = row[1] if isinstance(row, (list, tuple)) else row.get("contact_name")
            result["contact_name"] = contact if contact and str(contact).strip() not in ("", "None", "NOT KNOWN") else None
            return result
    except Exception:
        pass

    return result


def resolve_municipality_id(conn, municipality_name):
    """Look up a municipality's database ID from its display name.

    Tries exact match first, then case-insensitive LIKE match.

    Args:
        conn: Database connection (PgWrapper or sqlite3 connection)
        municipality_name: Display name (e.g., "Bayham", "Chatham-Kent")

    Returns:
        int or None: The municipality ID, or None if not found
    """
    def _execute(query, params):
        if hasattr(conn, "execute"):
            return conn.execute(query, params)
        raw = getattr(conn, "conn", conn)
        cur = raw.cursor()
        cur.execute(query, params)
        return cur

    # Exact match
    try:
        cur = _execute("SELECT id FROM municipalities WHERE name = ?", (municipality_name,))
        row = cur.fetchone()
        if row:
            return row[0] if isinstance(row, (list, tuple)) else row["id"]
    except Exception:
        pass

    # Case-insensitive fallback
    try:
        cur = _execute("SELECT id FROM municipalities WHERE LOWER(name) = LOWER(?)", (municipality_name,))
        row = cur.fetchone()
        if row:
            return row[0] if isinstance(row, (list, tuple)) else row["id"]
    except Exception:
        pass

    return None


# ══════════════════════════════════════════════════════════════
# Template Letter Filling
# ══════════════════════════════════════════════════════════════

# All known placeholders used across the 6 template letters
_PLACEHOLDER_MAP = {
    "[INSERT DATE]": "date",
    "[INSERT MUNICIPALITY NAME]": "municipality_name",
    "[INSERT ADDRESS]": "address",
    "[INSERT COUNTY FEDERATION NAME]": "county_federation",
    "[INSERT BYLAW NAME]": "bylaw_name",
    "[INSERT NAME]": "sender_name",
    "[INSERT TITLE]": "sender_title",
    "[INSERT CONTACT INFORMATION]": "contact_info",
}


def _replace_in_paragraph(paragraph, fields):
    """Replace [INSERT ...] placeholders in a single paragraph.

    Handles the tricky case where Word splits a placeholder across
    multiple runs (e.g., "[INSERT ", "NAME", "]" as three separate runs).
    """
    # First, try simple run-level replacement (covers most cases)
    for run in paragraph.runs:
        for placeholder, field_key in _PLACEHOLDER_MAP.items():
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, fields.get(field_key, ""))

    # Second pass: check if any placeholders span multiple runs
    full_text = paragraph.text
    for placeholder, field_key in _PLACEHOLDER_MAP.items():
        if placeholder in full_text:
            # Reconstruct the paragraph with the replacement
            # Merge all runs into the first run, then clear the rest
            combined = "".join(r.text for r in paragraph.runs)
            combined = combined.replace(placeholder, fields.get(field_key, ""))

            if paragraph.runs:
                paragraph.runs[0].text = combined
                for run in paragraph.runs[1:]:
                    run.text = ""


def fill_letter_template(template_path, fields):
    """Open a .docx template and replace all [INSERT ...] placeholders.

    Args:
        template_path: Absolute path to the template .docx file
        fields: dict with keys matching _PLACEHOLDER_MAP values:
            - date: str (e.g., "May 4, 2026")
            - municipality_name: str
            - address: str (can be empty)
            - county_federation: str
            - bylaw_name: str
            - sender_name: str
            - sender_title: str
            - contact_info: str (email/phone)

    Returns:
        BytesIO containing the personalized .docx document
    """
    doc = Document(template_path)

    # Replace in all paragraphs (body)
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, fields)
        # Bold the subject line automatically
        if paragraph.text.strip().startswith("Re:"):
            for run in paragraph.runs:
                run.bold = True

    # Replace in tables (some letters may use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, fields)

    # Replace in headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header and header.paragraphs:
                for paragraph in header.paragraphs:
                    _replace_in_paragraph(paragraph, fields)
        for footer in [section.footer, section.first_page_footer]:
            if footer and footer.paragraphs:
                for paragraph in footer.paragraphs:
                    _replace_in_paragraph(paragraph, fields)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def letter_to_plain_text(template_path, fields):
    """Render the personalized letter as plain text for email body.

    Reads the template, replaces placeholders, and extracts all paragraph
    text as a clean plain-text string.

    Args:
        template_path: Absolute path to the template .docx file
        fields: dict of placeholder values (same as fill_letter_template)

    Returns:
        str: The full letter as plain text
    """
    doc = Document(template_path)

    lines = []
    for paragraph in doc.paragraphs:
        text = paragraph.text
        for placeholder, field_key in _PLACEHOLDER_MAP.items():
            text = text.replace(placeholder, fields.get(field_key, ""))
        lines.append(text)

    # Join with newlines, collapse excessive blank lines
    raw = "\n".join(lines)
    # Collapse 3+ consecutive newlines into 2
    raw = re.sub(r"\n{3,}", "\n\n", raw)
    return raw.strip()


# ══════════════════════════════════════════════════════════════
# Mailto Link Generation
# ══════════════════════════════════════════════════════════════

def generate_mailto_link(recipient_email, subject, body_text):
    """Build a mailto: URL with pre-filled recipient, subject, and body.

    Handles URL encoding. If the total URL exceeds ~2000 characters
    (browser limit), the body is truncated with a continuation note.

    Args:
        recipient_email: str — the To: address
        subject: str — email subject line
        body_text: str — full letter as plain text

    Returns:
        str: A mailto: URL string
    """
    # mailto: URL limit is roughly 2000 chars in most browsers
    MAX_URL_LENGTH = 1900

    # Encode subject
    encoded_subject = urllib.parse.quote(subject, safe="")

    # Calculate space available for body
    base = f"mailto:{recipient_email}?subject={encoded_subject}&body="
    remaining = MAX_URL_LENGTH - len(base)

    # Encode body — if too long, truncate
    encoded_body = urllib.parse.quote(body_text, safe="")

    if len(encoded_body) > remaining:
        # Truncate the raw text and re-encode
        # Estimate: URL encoding roughly triples length for special chars
        safe_length = remaining // 3
        truncated = body_text[:safe_length].rsplit(" ", 1)[0]  # word boundary
        truncated += "\n\n[... Letter continues — please see attached .docx for full text ...]"
        encoded_body = urllib.parse.quote(truncated, safe="")

    return f"mailto:{recipient_email}?subject={encoded_subject}&body={encoded_body}"


def build_email_subject(municipality_name, bylaw_category_name):
    """Generate a professional email subject line.

    Args:
        municipality_name: str — e.g., "Township of Augusta"
        bylaw_category_name: str — e.g., "Development Charges"

    Returns:
        str: Subject line
    """
    return f"Request for Agricultural Exemption — {bylaw_category_name} — {municipality_name}"


def build_cover_email_body(sender_name, sender_title, county_federation, municipality_name, bylaw_category_name):
    """Generate a short cover email asking the clerk to review the attached letter.

    Args:
        sender_name: User's name
        sender_title: User's title
        county_federation: County federation name
        municipality_name: Target municipality
        bylaw_category_name: Target bylaw category

    Returns:
        str: Plain text cover email body
    """
    body = (
        f"Dear Clerk / Council of {municipality_name},\n\n"
        f"Please find attached a letter on behalf of {county_federation} regarding "
        f"agricultural exemptions under your {bylaw_category_name} bylaw.\n\n"
        f"We ask that this correspondence be included in the next council agenda or "
        f"directed to the appropriate staff for review.\n\n"
        f"Thank you for your time and consideration of this important matter for our local farm businesses.\n\n"
        f"Sincerely,\n\n"
        f"{sender_name}\n"
        f"{sender_title}\n"
        f"{county_federation}"
    )
    return body


def build_fields(sender_name, sender_title, county_federation, contact_info,
                 municipality_name, bylaw_name="", address=""):
    """Convenience function to build the fields dict for template filling.

    Args:
        sender_name: User's name
        sender_title: User's title/position
        county_federation: County Federation name
        contact_info: Email/phone string
        municipality_name: Target municipality
        bylaw_name: Name of the specific bylaw (optional)
        address: Municipality address (optional)

    Returns:
        dict: Ready for fill_letter_template() and letter_to_plain_text()
    """
    today = datetime.date.today()
    return {
        "date": today.strftime("%B %d, %Y"),
        "municipality_name": municipality_name or "",
        "address": address or "",
        "county_federation": county_federation or "",
        "bylaw_name": bylaw_name or "",
        "sender_name": sender_name or "",
        "sender_title": sender_title or "",
        "contact_info": contact_info or "",
    }


# ══════════════════════════════════════════════════════════════
# Advocacy Logging & Analytics
# ══════════════════════════════════════════════════════════════

def log_advocacy_action(conn, municipality_id, municipality_name,
                        bylaw_category, recipient_email,
                        sender_name, sender_org, action):
    """Log an advocacy action (letter download or mailto open) to the database.

    Args:
        conn: Database connection (PgWrapper or sqlite3)
        municipality_id: int or None
        municipality_name: str
        bylaw_category: str (e.g., "Development Charges")
        recipient_email: str or None
        sender_name: str
        sender_org: str (county federation name)
        action: str — 'letter_downloaded' or 'mailto_opened'
    """
    now = datetime.datetime.now().isoformat(timespec="seconds")
    try:
        # Ensure the table exists (safe for first run)
        if hasattr(conn, "execute"):
            conn.execute("""
                CREATE TABLE IF NOT EXISTS advocacy_log (
                    id              SERIAL PRIMARY KEY,
                    timestamp       TEXT NOT NULL,
                    municipality_id INTEGER,
                    municipality_name TEXT,
                    bylaw_category  TEXT,
                    recipient_email TEXT,
                    sender_name     TEXT,
                    sender_org      TEXT,
                    action          TEXT
                )
            """, ())
            conn.execute("""
                INSERT INTO advocacy_log
                    (timestamp, municipality_id, municipality_name,
                     bylaw_category, recipient_email, sender_name, sender_org, action)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, (now, municipality_id, municipality_name,
                  bylaw_category, recipient_email, sender_name, sender_org, action))
            if hasattr(conn, "commit"):
                conn.commit()
        else:
            raw = getattr(conn, "conn", conn)
            raw.execute("""
                CREATE TABLE IF NOT EXISTS advocacy_log (
                    id              INTEGER PRIMARY KEY AUTOINCREMENT,
                    timestamp       TEXT NOT NULL,
                    municipality_id INTEGER,
                    municipality_name TEXT,
                    bylaw_category  TEXT,
                    recipient_email TEXT,
                    sender_name     TEXT,
                    sender_org      TEXT,
                    action          TEXT
                )
            """)
            raw.execute("""
                INSERT INTO advocacy_log
                    (timestamp, municipality_id, municipality_name,
                     bylaw_category, recipient_email, sender_name, sender_org, action)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, (now, municipality_id, municipality_name,
                  bylaw_category, recipient_email, sender_name, sender_org, action))
            raw.commit()
    except Exception:
        pass  # Don't break the UI if logging fails


def get_advocacy_stats(conn):
    """Retrieve aggregate advocacy usage statistics.

    Returns:
        dict with keys:
            total_actions: int — total logged actions
            total_downloads: int — letter downloads
            total_emails: int — mailto opens
            municipalities_contacted: int — unique municipalities
            by_municipality: list of dicts — [{name, count}] sorted desc
            by_category: list of dicts — [{category, count}] sorted desc
            recent: list of dicts — last 10 actions
    """
    stats = {
        "total_actions": 0,
        "total_downloads": 0,
        "total_emails": 0,
        "municipalities_contacted": 0,
        "by_municipality": [],
        "by_category": [],
        "recent": [],
    }

    def _execute(query, params=()):
        if hasattr(conn, "execute"):
            return conn.execute(query, params)
        raw = getattr(conn, "conn", conn)
        cur = raw.cursor()
        cur.execute(query, params)
        return cur

    try:
        # Total counts
        cur = _execute("SELECT COUNT(*) FROM advocacy_log")
        row = cur.fetchone()
        stats["total_actions"] = row[0] if row else 0

        cur = _execute("SELECT COUNT(*) FROM advocacy_log WHERE action = ?",
                       ("letter_downloaded",))
        row = cur.fetchone()
        stats["total_downloads"] = row[0] if row else 0

        cur = _execute("SELECT COUNT(*) FROM advocacy_log WHERE action = ?",
                       ("mailto_opened",))
        row = cur.fetchone()
        stats["total_emails"] = row[0] if row else 0

        # Unique municipalities
        cur = _execute("SELECT COUNT(DISTINCT municipality_name) FROM advocacy_log")
        row = cur.fetchone()
        stats["municipalities_contacted"] = row[0] if row else 0

        # By municipality
        cur = _execute("""
            SELECT municipality_name, COUNT(*) as cnt
            FROM advocacy_log
            GROUP BY municipality_name
            ORDER BY cnt DESC
        """)
        stats["by_municipality"] = [
            {"name": r[0], "count": r[1]} for r in cur.fetchall()
        ]

        # By category
        cur = _execute("""
            SELECT bylaw_category, COUNT(*) as cnt
            FROM advocacy_log
            GROUP BY bylaw_category
            ORDER BY cnt DESC
        """)
        stats["by_category"] = [
            {"category": r[0], "count": r[1]} for r in cur.fetchall()
        ]

        # Recent 10
        cur = _execute("""
            SELECT timestamp, municipality_name, bylaw_category,
                   sender_name, sender_org, action
            FROM advocacy_log
            ORDER BY timestamp DESC
            LIMIT 10
        """)
        stats["recent"] = [
            {"timestamp": r[0], "municipality": r[1], "category": r[2],
             "sender": r[3], "org": r[4], "action": r[5]}
            for r in cur.fetchall()
        ]
    except Exception:
        pass  # Table may not exist yet

    return stats
