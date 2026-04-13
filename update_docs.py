import docx
import datetime

filepath = r'docs/Municipal_Bylaw_Database_Developer_Manual.docx'
doc = docx.Document(filepath)

doc.add_heading('Architecture Update: Neon Serverless PostgreSQL Migration', level=1)
doc.add_paragraph('Date: April 2, 2026')
doc.add_paragraph('Due to data persistence failures arising from Streamlit Cloud\'s containerized ephemeral filesystem, the system architecture has been permanently migrated away from the legacy local SQLite database (bylaws.db).')
doc.add_heading('Key Upgrades', level=2)
doc.add_paragraph('1. Database Engine: The native stack is now connected directly to a Neon.tech Serverless PostgreSQL cluster. The local bylaws.db file is no longer used by the production environment.', style='List Bullet')
doc.add_paragraph('2. psycopg2 Python Integration: We replaced standard sqlite3 logic with psycopg2. Because the Streamlit DataFrame library native output relies heavily on Numpy, global register_adapter rules have been explicitly hardcoded into db_utils.py to ensure psycopg2 doesn\'t throw exceptions when rendering np.int64 variables.', style='List Bullet')
doc.add_paragraph('3. Query Parameter Interpolation: The legacy SQLite wildcard marker (?) has been deprecated system-wide. All backend and GUI queries rely strictly upon strict string formatting (%s) mapped to standard Postgres parameter parsing strings.', style='List Bullet')
doc.add_heading('Future Development Guidelines', level=2)
p1 = doc.add_paragraph('If you are writing additional reports, custom data scrapers, or altering the dashboard: ')
p1.add_run('Always format parameterized variables with %s').bold = True
p1.add_run(' instead of ?, and map bracketed columns like [exemption_status] strictly into strings surrounded by double quotes (e.g. "exemption_status").')

doc.save(filepath)
print('Document updated successfully!')
