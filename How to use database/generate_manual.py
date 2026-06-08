"""
Generate the OFA Municipal Bylaw Database onboarding manual as a Word document.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Style adjustments ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Heading styles
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)  # Dark green

# ══════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Ontario Municipal Bylaw Database')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('User Onboarding Guide')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()

org = doc.add_paragraph()
org.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = org.add_run('Ontario Federation of Agriculture')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run('June 2026 — Version 1.0')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    ('1.', 'What Is This Tool?'),
    ('2.', 'Getting Started'),
    ('3.', 'The Home Page'),
    ('4.', 'Lower Tier Map'),
    ('5.', 'Upper Tier Map'),
    ('6.', 'Data Browser'),
    ('7.', 'Admin Panel'),
    ('8.', 'Intelligence Scanner'),
    ('9.', 'Report Generator'),
    ('10.', 'The Advocacy Letter Tool'),
    ('11.', 'Understanding the Colour Coding'),
    ('12.', 'Frequently Asked Questions'),
    ('13.', 'Glossary of Terms'),
]
for num, label in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  {label}')
    run.font.size = Pt(12)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# HELPER
# ══════════════════════════════════════════════════════════════
def add_tip(doc, text):
    """Add a highlighted tip box."""
    p = doc.add_paragraph()
    run = p.add_run('💡 Tip: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    run2 = p.add_run(text)
    run2.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

def add_note(doc, text):
    """Add a note box."""
    p = doc.add_paragraph()
    run = p.add_run('📌 Note: ')
    run.bold = True
    run2 = p.add_run(text)

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f' — {text}')
    else:
        p.add_run(text)


# ══════════════════════════════════════════════════════════════
# SECTION 1 — WHAT IS THIS TOOL?
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. What Is This Tool?', level=1)

doc.add_paragraph(
    'The Ontario Municipal Bylaw Database is a web-based application built by the Ontario Federation '
    'of Agriculture (OFA). It tracks municipal bylaws across all 444 municipalities in Ontario and '
    'helps OFA staff, board members, and county federations understand how local bylaws affect farming operations.'
)

doc.add_paragraph(
    'The tool covers seven different bylaw categories:'
)

categories = [
    ('Development Charges (DC)', 'Fees charged to developers for new growth. Farms may or may not be exempt.'),
    ('Stormwater', 'Stormwater management fees. Some municipalities charge farms; others exempt them.'),
    ('Site Alteration & Fill', 'Rules about moving soil, grading land, or accepting fill. Farms often have exemptions.'),
    ('Livestock Guardian Dogs (LGD)', 'Municipal animal control bylaws that may or may not recognize livestock guardian dogs and herding dogs as working farm animals.'),
    ('Tree / Forest Conservation', 'Bylaws that regulate tree cutting. Many have exemptions for normal farming.'),
    ('Backyard Chickens', 'Whether residents can keep chickens. Tracks limits, licencing, and rooster rules.'),
    ('Fence Bylaws', 'Municipal fence bylaws, including whether they replace the provincial Line Fences Act and allow electric or security fencing for farms.'),
]
for name, desc in categories:
    add_bullet(doc, desc, name)

doc.add_paragraph(
    'For each municipality and category, the database tracks whether a farm exemption exists, '
    'the bylaw name and link, when it was enacted, when it expires, and detailed category-specific '
    'information.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 2 — GETTING STARTED
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. Getting Started', level=1)

doc.add_heading('How to Access the Tool', level=2)
doc.add_paragraph(
    'The database is hosted on Streamlit Community Cloud and can be accessed through any modern '
    'web browser (Chrome, Edge, Firefox, Safari). You will be given a URL link — simply click it '
    'or paste it into your browser\'s address bar.'
)

doc.add_heading('Navigating the Application', level=2)
doc.add_paragraph(
    'The application has a sidebar on the left side of the screen. This sidebar serves two purposes:'
)
add_bullet(doc, 'Navigation — Click on any page name in the sidebar to switch between pages.')
add_bullet(doc, 'Filters and controls — Each page has its own set of filters that appear in the sidebar when you are on that page.')

doc.add_paragraph(
    'The available pages are listed in the sidebar:'
)
pages = [
    ('Home', 'Welcome page with an overview of the tool.'),
    ('Lower Tier Map', 'Interactive map of all 414 local (lower-tier) municipalities.'),
    ('Upper Tier Map', 'Interactive map of all 33 regional (upper-tier) municipalities.'),
    ('Data Browser', 'Searchable table view of all bylaw data.'),
    ('Admin', 'Password-protected page for editing database records (staff only).'),
    ('Intelligence Scanner', 'Tool to scan municipal websites for bylaw-related keywords.'),
    ('Report Generator', 'Create professional Word document reports.'),
]
for name, desc in pages:
    add_bullet(doc, desc, name)

add_tip(doc, 'If the sidebar is hidden, click the ">" arrow in the top-left corner of the screen to open it.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 3 — HOME PAGE
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. The Home Page', level=1)

doc.add_paragraph(
    'When you first open the application, you land on the Home page. This page provides an '
    'overview of what the tool does and quick links to get started.'
)

doc.add_heading('What You Will See', level=2)
add_bullet(doc, 'A welcome message explaining the purpose of the database.')
add_bullet(doc, 'A navigation table describing the two map views (Lower Tier and Upper Tier).')
add_bullet(doc, 'An explanation of the Intelligence Engine — the automated weekly scanner that checks municipal websites for bylaw changes.')
add_bullet(doc, 'Quick links to jump directly to the Intelligence Scanner, signal filters, expiry alerts, and map features.')

add_tip(doc, 'Use the Home page as a starting point. Read through the overview once to understand the tool, then use the sidebar to navigate to the page you need.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 4 — LOWER TIER MAP
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. Lower Tier Map', level=1)

doc.add_paragraph(
    'The Lower Tier Map is the primary visual tool in the database. It shows an interactive map of '
    'all 414 lower-tier (local) municipalities in Ontario, colour-coded by their bylaw exemption status.'
)

doc.add_heading('How to Use the Map', level=2)

doc.add_heading('Step 1: Choose a Bylaw Category', level=3)
doc.add_paragraph(
    'In the sidebar, use the "Bylaw" dropdown to select which bylaw category you want to view. '
    'For example, selecting "Farm Exemption for Development Charges" will colour the map based on '
    'whether each municipality exempts farms from development charges.'
)

doc.add_heading('Step 2: Apply Filters (Optional)', level=3)
doc.add_paragraph('You can narrow down what the map shows using the sidebar filters:')
add_bullet(doc, 'Shows only municipalities with a specific status (YES, NO, or N/A).', 'Status filter')
add_bullet(doc, 'Highlights municipalities where bylaws are expiring soon, already expired, or have unknown expiry dates.', 'Expiry alert filter')
add_bullet(doc, 'When checked, the map highlights only municipalities where our weekly scanner has found new bylaw-related activity on their website (e.g., if a town recently posted an agenda mentioning a "development charges update", the scanner picks it up and flags the town on the map).', 'Scanner Signals checkbox')
add_bullet(doc, 'Choose "Highlight matches" to dim non-matching municipalities, or "Filter to matches" to hide them entirely.', 'Display mode')
add_bullet(doc, 'Type part of a municipality name to find it quickly.', 'Search')

doc.add_heading('Step 3: Read the Map', level=3)
doc.add_paragraph('The map uses colours to show the status of each municipality:')
add_bullet(doc, 'The municipality has a farm exemption (YES).', 'Green')
add_bullet(doc, 'The municipality does NOT have a farm exemption (NO).', 'Red')
add_bullet(doc, 'Information is not available or not applicable (N/A).', 'Grey')

doc.add_paragraph('The map also uses border outlines to show additional information:')
add_bullet(doc, 'Our scanner recently found bylaw-related activity on this municipality\'s website.', 'Red-Orange border')
add_bullet(doc, 'The bylaw is expiring within the next year.', 'Orange border')
add_bullet(doc, 'The bylaw has already expired.', 'Black border')
add_bullet(doc, 'The bylaw\'s expiry date is unknown.', 'Yellow border')

doc.add_heading('Step 4: Click on a Municipality', level=3)
doc.add_paragraph(
    'Click on any municipality on the map to select it. When you click, the sidebar will update to show '
    'detailed information about that municipality, including:'
)
add_bullet(doc, 'The bylaw name and its current status.')
add_bullet(doc, 'The date the bylaw was enacted and when it expires.')
add_bullet(doc, 'A link to view the actual bylaw document (if available).')
add_bullet(doc, 'Any scanner intelligence signals found for that municipality.')

doc.add_heading('Step 5: View Summary Metrics', level=3)
doc.add_paragraph(
    'At the top of the page, above the map, you will see a row of summary numbers showing how many '
    'municipalities fall into each status category (YES, NO, N/A, Expiring Soon, etc.). These numbers '
    'update as you apply filters.'
)

doc.add_heading('Step 6: Advocacy Tools and Policy Positions', level=3)
doc.add_paragraph(
    'Scroll down below the map to find two additional resources for any selected municipality:'
)
add_bullet(doc, 'OFA\'s official policy position and background information on the specific bylaw topic you are viewing.', 'OFA Position Statement')

doc.add_heading('How to Generate an Advocacy Letter', level=3)
doc.add_paragraph(
    'If you select a municipality that does NOT currently have a farm exemption (i.e., its status is "NO"), '
    'the database will activate the Advocacy Letter Generator below the map. This feature allows you to instantly '
    'draft a customized letter to local officials. Here is how to use it:'
)
add_bullet(doc, 'Scroll down below the main map window and click the collapsible box titled "✉️ Send Personalized Advocacy Letter".')
add_bullet(doc, 'Fill in the text boxes provided (e.g., the local representative\'s name, your name, your title, and your contact information).')
add_bullet(doc, 'Click the "👁️ Preview Letter" button to read the letter with your details filled in.')
add_bullet(doc, 'Click the "📥 Download Personalized Letter" button to save a perfectly formatted Microsoft Word (.docx) file to your computer.')
add_bullet(doc, 'If an email address is on file for that municipality, you can also click the "📧 Generate Letter & Open Email Draft" button. This will automatically open your computer\'s default email program (like Outlook) with a pre-written message, subject line, and the recipient\'s email address already filled in! All you have to do is attach the .docx file you downloaded and hit send.')

doc.add_heading('Exporting Map Data', level=2)
doc.add_paragraph(
    'At the bottom of the page, expand the "Results (filtered)" section to see a table of all the '
    'municipalities currently shown on the map. You can download this table as a CSV file by clicking '
    'the download button.'
)

add_tip(doc, 'Hover your mouse over a municipality on the map to see a quick tooltip with its name, status, and bylaw details without clicking.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 5 — UPPER TIER MAP
# ══════════════════════════════════════════════════════════════
doc.add_heading('5. Upper Tier Map', level=1)

doc.add_paragraph(
    'The Upper Tier Map works the same way as the Lower Tier Map, but it shows Ontario\'s 33 '
    'upper-tier (regional) municipalities instead of the 414 local municipalities.'
)

doc.add_paragraph(
    'Upper-tier municipalities include regional governments like the Region of Peel, Region of '
    'Waterloo, County of Oxford, and similar entities. Some bylaw categories (like Development '
    'Charges) can apply at both the lower-tier and upper-tier level.'
)

doc.add_paragraph(
    'All the same filters, colour coding, municipality detail panel, and export features from the '
    'Lower Tier Map are available here. The only difference is the set of municipalities shown.'
)

add_tip(doc, 'Use the Upper Tier Map when you want to see the "big picture" at the regional level. Use the Lower Tier Map for municipality-by-municipality detail.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 6 — DATA BROWSER
# ══════════════════════════════════════════════════════════════
doc.add_heading('6. Data Browser', level=1)

doc.add_paragraph(
    'The Data Browser is a detailed table view of all the bylaw data in the database. While the maps '
    'give you a visual overview, the Data Browser lets you dig into the specifics — search, filter, '
    'sort, and read the detailed bylaw information for every municipality.'
)

doc.add_heading('Choosing a Bylaw Category', level=2)
doc.add_paragraph(
    'Use the "Bylaw Category" dropdown in the sidebar to select which type of bylaw you want to browse. '
    'The seven options are: Development Charges, Stormwater, Site Alteration & Fill, Livestock Guardian Dogs, '
    'Tree / Forest Conservation, Backyard Chickens, and Fence Bylaws.'
)

doc.add_heading('Filtering the Data', level=2)
doc.add_paragraph('The sidebar provides several ways to narrow down the data:')
add_bullet(doc, 'Type part of a municipality name to find it in the table.', 'Search municipality')
add_bullet(doc, 'Filter by county or region (e.g., show only municipalities in Oxford County).', 'Geographic Area')
add_bullet(doc, 'Filter by Yes, No, N/A, or NOT KNOWN.', 'Exemption Status')

doc.add_heading('Reading the Data Table', level=2)
doc.add_paragraph(
    'The main table shows one row per municipality. The columns include:'
)
add_bullet(doc, 'The name of the municipality.', 'Municipality')
add_bullet(doc, 'Whether it is Lower Tier, Upper Tier, or Single Tier.', 'Tier')
add_bullet(doc, 'The county or region the municipality belongs to.', 'County / Region')
add_bullet(doc, 'Whether a farm exemption exists (colour-coded green, red, or yellow).', 'Farm Exemption')
add_bullet(doc, 'The official name of the bylaw.', 'Bylaw Name')
add_bullet(doc, 'When the bylaw was enacted and when it expires.', 'Enacted / Expires')
add_bullet(doc, 'Notes about the bylaw\'s expiry.', 'Expiry Notes')
add_bullet(doc, 'The review progress status (Complete, In Progress, etc.).', 'Progress')
add_bullet(doc, 'A clickable link to view the actual bylaw document online (if available).', 'Bylaw Link')

doc.add_paragraph(
    'Depending on the bylaw category you selected, additional columns will appear. For example, if you '
    'selected "Livestock Guardian Dogs," you will see extra columns for LGD Definition, Herding Definition, '
    'License Fee Exemption, Collar/Tag Requirements, and more.'
)

doc.add_heading('Municipality Detail Panel', level=2)
doc.add_paragraph(
    'Below the table, you will find the Municipality Detail section. Select any municipality from the '
    'dropdown to see its complete profile, including:'
)
add_bullet(doc, 'Basic information — name, tier, region, zone, website.')
add_bullet(doc, 'Contact information — contact name, position, email, and phone number.')
add_bullet(doc, 'All seven bylaw categories in expandable sections (click to expand/collapse each one).')
add_bullet(doc, 'AMCTO Directory contacts — a table of municipal staff from the AMCTO directory.')
add_bullet(doc, 'Scanner Intelligence Signals — any recent scanner hits for this municipality.')

add_tip(doc, 'Use the "Export" section in the sidebar to download the current filtered view as a CSV file that you can open in Excel.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 7 — ADMIN PANEL
# ══════════════════════════════════════════════════════════════
doc.add_heading('7. Admin Panel', level=1)

doc.add_paragraph(
    'The Admin Panel is a password-protected page used by OFA staff to update and manage the database. '
    'Only authorized users should use this page.'
)

doc.add_heading('Logging In', level=2)
doc.add_paragraph(
    'When you navigate to the Admin page, you will see a password field. Enter the admin password and click '
    '"Login." If you have forgotten the password, click the "Forgot Password?" link to reset it to the default.'
)

doc.add_heading('Three Modes', level=2)
doc.add_paragraph('Once logged in, use the sidebar to choose between three modes:')

doc.add_heading('Mode 1: Edit Municipality', level=3)
doc.add_paragraph(
    'This is where you update bylaw data for individual municipalities. Here is how to use it:'
)
doc.add_paragraph('1. Select a municipality from the dropdown.')
doc.add_paragraph('2. You will see a row of tabs across the top — one for "Municipality Info" and one for each of the seven bylaw categories.')
doc.add_paragraph('3. Click on the tab you want to edit.')
doc.add_paragraph('4. Fill in or update the fields in the form.')
doc.add_paragraph('5. Click the "Save" button at the bottom of the form to save your changes.')

doc.add_paragraph('The "Municipality Info" tab lets you update:')
add_bullet(doc, 'Municipality name, tier status, geographic area, and zone.')
add_bullet(doc, 'Website URL.')
add_bullet(doc, 'Contact name, position, email, and phone.')

doc.add_paragraph('Each bylaw category tab lets you update:')
add_bullet(doc, 'Bylaw name and link to the bylaw document.')
add_bullet(doc, 'Date enacted and expiry date.')
add_bullet(doc, 'Farm exemption status (Yes / No / NOT KNOWN / N/A / No explicit exemption found).')
add_bullet(doc, 'Exemption wording and other notes.')
add_bullet(doc, 'Progress status (e.g., Complete, In Progress).')
add_bullet(doc, 'Category-specific detail fields (varies by category).')

doc.add_heading('Mode 2: Change History', level=3)
doc.add_paragraph(
    'This mode shows an audit log of every edit made to the database. You can see who made changes, what '
    'was changed, the old value, and the new value. Use the slider to control how many entries to show, and '
    'download the history as a CSV file.'
)

doc.add_heading('Mode 3: Bulk Import', level=3)
doc.add_paragraph(
    'This mode allows you to upload a CSV file to update multiple records at once. The CSV must have columns '
    'for municipality_name, category, field, and value. This is useful for making large batches of updates '
    'without editing each municipality one by one.'
)

doc.add_heading('Syncing Changes to the Maps', level=2)
doc.add_paragraph(
    'After making edits in the Admin panel, click the "Sync Edits to Maps" button in the sidebar. '
    'This rebuilds the map data files so that your changes appear on the Lower Tier and Upper Tier maps. '
    'If you do not sync, the maps may show outdated information until the next automatic sync.'
)

add_note(doc, 'All changes are logged automatically. Every edit is recorded in the Change History with a timestamp, so you can always see what was changed and when.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 8 — INTELLIGENCE SCANNER
# ══════════════════════════════════════════════════════════════
doc.add_heading('8. Intelligence Scanner', level=1)

doc.add_paragraph(
    'The Intelligence Scanner is a powerful tool that automatically searches municipal websites for '
    'keywords related to bylaw changes. It helps OFA stay informed about new bylaws, bylaw reviews, '
    'and upcoming council decisions that could affect agriculture.'
)

doc.add_heading('The Three Tabs', level=2)

doc.add_heading('Tab 1: Live Target Scanner', level=3)
doc.add_paragraph(
    'This tab lets you run a real-time scan of municipal websites right now. Here is how to use it:'
)
doc.add_paragraph('1. Choose your keywords. You can select from preset keyword packs (e.g., "ALTO Rail," "Plant-Based Treaty," "Ontario Foodbelt") or type your own custom keywords, one per line.')
doc.add_paragraph('2. Choose which municipalities to scan. You can select "Select All" to scan all 444 municipalities (this takes 3-5 minutes) or pick specific municipalities from the dropdown.')
doc.add_paragraph('3. Click "Launch Scan" to start the scan.')
doc.add_paragraph('4. A progress bar will show you how the scan is going. You will see notifications as hits are found.')
doc.add_paragraph('5. When the scan finishes, you will see:')
add_bullet(doc, 'A coverage report showing how many municipalities were successfully scanned, how many documents were read, and any errors.')
add_bullet(doc, 'A list of keyword hits showing which municipalities had matches, what keyword was found, a snippet of the surrounding text, and a link to the source document.')
doc.add_paragraph('6. Download the full results as a CSV file for further analysis or sharing.')

doc.add_heading('Tab 2: Historical Intelligence Database', level=3)
doc.add_paragraph(
    'This tab acts as the master "Library" for the 7 official OFA advocacy categories '
    '(like Development Charges, Stormwater Fees, and Backyard Chickens) across all 444 municipalities. '
    'Rather than making you wait for a live scan, this tab relies on an automated script '
    'that quietly runs every week in the background, reading thousands of documents to find '
    'OFA-specific issues, and saving the results here. You can filter by county/region, '
    'bylaw category, or search for specific keywords inside the extracted text snippets. '
    'This is your province-wide early warning system for tracking when local councils start '
    'quietly laying the groundwork for new agricultural bylaws.'
)

doc.add_heading('Tab 3: Portal Health Monitor', level=3)
doc.add_paragraph(
    'This tab identifies municipalities whose websites may be out of date or inaccessible. If a '
    'municipality\'s portal has not had a readable document in 60 or more days, it shows up here as "stale." '
    'This helps OFA know which municipalities might need to be checked manually.'
)

add_tip(doc, 'Use the Live Scanner when you need to check for something specific (like a new OFA advocacy topic). Use the Historical Database to review what the automated weekly scans have found over time.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 9 — REPORT GENERATOR
# ══════════════════════════════════════════════════════════════
doc.add_heading('9. Report Generator', level=1)

doc.add_paragraph(
    'The Report Generator creates professional Word documents (.docx) that summarize bylaw data. '
    'These reports are formatted with OFA branding and can be shared with board members, county '
    'federations, or other stakeholders.'
)

doc.add_heading('How to Generate a Report', level=2)
doc.add_paragraph('1. Choose your report scope:')
add_bullet(doc, 'Covers all 444 municipalities across Ontario. Best for province-wide overviews.', 'Provincial')
add_bullet(doc, 'Covers all municipalities within a specific county or region. Best for county-level meetings.', 'County / Region')
add_bullet(doc, 'Covers a single municipality in depth. Best for preparing for a specific meeting or advocacy effort.', 'Municipality')

doc.add_paragraph('2. If you selected County or Municipality, choose the specific one from the dropdown that appears.')
doc.add_paragraph('3. Click "Generate Report." The tool will create the Word document.')
doc.add_paragraph('4. Click "Download Report (.docx)" to save it to your computer.')

doc.add_heading('What Is in the Report', level=2)
doc.add_paragraph('Depending on the scope you choose, the report may include:')
add_bullet(doc, 'A high-level summary of findings.', 'Executive Summary')
add_bullet(doc, 'A quick-reference table showing exemption counts across all seven categories.', 'Exemption Scorecard')
add_bullet(doc, 'Detailed tables for each of the seven bylaw types.', 'Category Sections')
add_bullet(doc, 'Individual profiles with complete bylaw details (County and Municipality reports).', 'Municipality Profiles')
add_bullet(doc, 'Recent scanner hits and alerts about bylaw changes.', 'Intelligence Alerts')
add_bullet(doc, 'A ranked list of municipalities that are priorities for advocacy outreach (Provincial reports).', 'Advocacy Priority Matrix')
add_bullet(doc, 'An explanation of how data was collected and what the status codes mean.', 'Methodology')

add_tip(doc, 'Generate a County report before a county federation meeting to give attendees a summary of all the bylaws in their area.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 10 — ADVOCACY LETTER TOOL
# ══════════════════════════════════════════════════════════════
doc.add_heading('10. The Advocacy Letter Tool', level=1)

doc.add_paragraph(
    'The map pages (both Lower Tier and Upper Tier) include a built-in advocacy letter tool. '
    'This feature appears below the map when you are viewing a municipality that does NOT have '
    'a farm exemption (status = NO or NOT KNOWN).'
)

doc.add_heading('How to Use It', level=2)
doc.add_paragraph('1. On the map page, look for the section titled "Send Personalized Advocacy Letter." Click to expand it.')
doc.add_paragraph('2. Fill in your details:')
add_bullet(doc, 'Your Name (required)')
add_bullet(doc, 'Title / Position (required)')
add_bullet(doc, 'Organization / Affiliation (required)')
add_bullet(doc, 'Contact Email (optional)')
add_bullet(doc, 'Contact Phone (optional)')

doc.add_paragraph('3. The municipality name, bylaw topic, and recipient email are filled in automatically from the database.')
doc.add_paragraph('4. Click "Preview Letter" to see what the letter will say before sending.')
doc.add_paragraph('5. You have two options:')
add_bullet(doc, 'Downloads a Word document (.docx) with your personalized letter that you can print or email as an attachment.', 'Download Personalized Letter')
add_bullet(doc, 'Downloads the letter AND opens your email program with a draft email pre-addressed to the municipal clerk.', 'Generate Letter & Open Email Draft')

doc.add_heading('OFA Policy Positions', level=2)
doc.add_paragraph(
    'Below the map on each page, you will also see the OFA\'s official policy position on the selected '
    'bylaw topic. This gives you the background and talking points you need for advocacy conversations.'
)

doc.add_heading('Tracking Advocacy Activity', level=2)
doc.add_paragraph(
    'The tool automatically tracks how many letters have been downloaded and emails opened. '
    'You can see this data in the "Advocacy Tool Usage" section, which shows metrics like total '
    'letters generated, emails opened, and a list of municipalities contacted.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 11 — COLOUR CODING
# ══════════════════════════════════════════════════════════════
doc.add_heading('11. Understanding the Colour Coding', level=1)

doc.add_paragraph(
    'The application uses consistent colour coding throughout the maps and tables to help you '
    'quickly understand the status of each municipality.'
)

doc.add_heading('Map Fill Colours', level=2)
# Create a table for map colours
table = doc.add_table(rows=4, cols=2)
table.style = 'Light List Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Colour'
hdr[1].text = 'Meaning'
data = [
    ('Green', 'Farm exemption exists (YES)'),
    ('Red', 'No farm exemption (NO)'),
    ('Grey', 'Information not available or not applicable (N/A)'),
]
for i, (colour, meaning) in enumerate(data):
    row = table.rows[i + 1].cells
    row[0].text = colour
    row[1].text = meaning

doc.add_paragraph()
doc.add_heading('Map Border Outlines', level=2)
table2 = doc.add_table(rows=5, cols=2)
table2.style = 'Light List Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Border Colour'
hdr2[1].text = 'Meaning'
data2 = [
    ('Red-Orange', 'Scanner found recent bylaw activity on this municipality\'s website'),
    ('Orange', 'The bylaw is expiring soon (within the next year)'),
    ('Black', 'The bylaw has already expired'),
    ('Yellow', 'The bylaw\'s expiry date is unknown'),
]
for i, (colour, meaning) in enumerate(data2):
    row2 = table2.rows[i + 1].cells
    row2[0].text = colour
    row2[1].text = meaning

doc.add_paragraph()
doc.add_heading('Table Cell Colours (Data Browser)', level=2)
table3 = doc.add_table(rows=4, cols=2)
table3.style = 'Light List Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Cell Colour'
hdr3[1].text = 'Meaning'
data3 = [
    ('Green background', 'Yes — exemption exists'),
    ('Red background', 'No — no exemption'),
    ('Yellow background', 'N/A or NOT KNOWN'),
]
for i, (colour, meaning) in enumerate(data3):
    row3 = table3.rows[i + 1].cells
    row3[0].text = colour
    row3[1].text = meaning

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 12 — FAQ
# ══════════════════════════════════════════════════════════════
doc.add_heading('12. Frequently Asked Questions', level=1)

faqs = [
    ('How often is the data updated?',
     'The database is updated by OFA staff as new information becomes available. The Intelligence Scanner runs an automated weekly scan of municipal websites to flag potential changes. Staff then review the flagged items and update the database manually through the Admin panel.'),
    ('Can I break anything by clicking around?',
     'No. The maps, Data Browser, Intelligence Scanner, and Report Generator are all read-only for regular users. Only the Admin panel (which is password-protected) can make changes to the data.'),
    ('What does "NO BYLAW IN PLACE" mean in the Progress column?',
     'This means the municipality does not currently have a bylaw of that type. For example, a municipality with "NO BYLAW IN PLACE" for Stormwater means they do not charge stormwater fees.'),
    ('What does the scanner actually do?',
     'The scanner visits municipal websites and looks through council agendas, minutes, and posted documents for keywords related to each bylaw category (e.g., "Development Charges Background Study," "site alteration bylaw"). When it finds a match, it records the keyword, a snippet of surrounding text, and a link to the source document.'),
    ('Can I use this tool on my phone?',
     'The tool works best on a laptop or desktop computer with a large screen. The maps and tables are designed for wide screens. While it will technically work on a phone or tablet, the experience will be limited.'),
    ('Who do I contact if I find an error in the data?',
     'Contact OFA staff. They can update the data through the Admin panel. All corrections are logged in the Change History.'),
    ('How do I print or share a report?',
     'Use the Report Generator page to create a Word document, then download it. You can print it, email it as an attachment, or share it through OneDrive or other file-sharing tools.'),
    ('What is the difference between the Lower Tier and Upper Tier maps?',
     'Lower Tier municipalities (414) are the local, town-level governments. Upper Tier municipalities (33) are the regional or county-level governments. Some bylaws (like Development Charges) can exist at both levels, so a farm could be subject to bylaws from both their lower-tier and upper-tier municipality.'),
]

for question, answer in faqs:
    p = doc.add_paragraph()
    run = p.add_run(f'Q: {question}')
    run.bold = True
    doc.add_paragraph(f'A: {answer}')
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 13 — GLOSSARY
# ══════════════════════════════════════════════════════════════
doc.add_heading('13. Glossary of Terms', level=1)

glossary = [
    ('AMCTO', 'Association of Municipal Managers, Clerks and Treasurers of Ontario. A professional directory of municipal staff.'),
    ('Bylaw', 'A law passed by a municipal council that applies within that municipality\'s boundaries.'),
    ('CSV', 'Comma Separated Values. A simple file format that can be opened in Excel or Google Sheets.'),
    ('Development Charges (DC)', 'Fees that municipalities charge to fund infrastructure needed for new growth (roads, water, sewers, etc.).'),
    ('Exemption', 'A rule that says a particular group (e.g., farms) does not have to follow a specific part of a bylaw or does not have to pay a specific charge.'),
    ('Farm Exemption', 'A provision in a municipal bylaw that exempts agricultural operations from some or all requirements of that bylaw.'),
    ('GeoJSON', 'A file format used to represent geographic boundaries on maps.'),
    ('Intelligence Scanner', 'The automated tool that searches municipal websites for keywords related to bylaw changes.'),
    ('LGD', 'Livestock Guardian Dog. A dog used to protect livestock from predators.'),
    ('Lower Tier', 'A local (town, township, city, or village) municipality. Ontario has 414 lower-tier municipalities.'),
    ('N/A', 'Not Applicable. Used when a bylaw category does not apply to a particular municipality.'),
    ('NOT KNOWN', 'The information has not yet been researched or confirmed for this municipality.'),
    ('OFA', 'Ontario Federation of Agriculture.'),
    ('Parquet', 'A data file format used internally by the application to store map data efficiently.'),
    ('Scanner Signal', 'A "hit" or match found by the Intelligence Scanner when a keyword appears on a municipal website.'),
    ('Signal', 'See Scanner Signal.'),
    ('Single Tier', 'A municipality that functions as both a lower-tier and upper-tier government (e.g., City of Toronto, City of Hamilton).'),
    ('Site Alteration', 'The process of changing land — grading, filling, excavating, or removing topsoil.'),
    ('Stormwater', 'Rainwater and melted snow that runs off surfaces. Some municipalities charge fees to manage stormwater infrastructure.'),
    ('Upper Tier', 'A regional or county-level municipality. Ontario has 33 upper-tier municipalities.'),
]

table = doc.add_table(rows=len(glossary) + 1, cols=2)
table.style = 'Light List Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Term'
hdr[1].text = 'Definition'
for i, (term, definition) in enumerate(glossary):
    row = table.rows[i + 1].cells
    row[0].text = term
    p = row[0].paragraphs[0]
    p.runs[0].bold = True
    row[1].text = definition

# Set column widths
for row in table.rows:
    row.cells[0].width = Inches(1.8)
    row.cells[1].width = Inches(4.7)

doc.add_paragraph()
doc.add_paragraph()

# ── Final footer ──
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— End of Guide —')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(10)
run.italic = True

doc.add_paragraph()
footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = footer2.add_run('Ontario Federation of Agriculture © 2026')
run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run2.font.size = Pt(9)

# ── Save ──
output_dir = r"c:\Users\ben.lefort\OneDrive - Ontario Federation of Agriculture\Desktop\Municipal Bylaw Database\Data Pulls\Reports\Province Wide\All bylaws\county-bylaw-maps\How to use database"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "OFA Bylaw Database - User Onboarding Guide.docx")
doc.save(output_path)
print(f"Manual saved to: {output_path}")
